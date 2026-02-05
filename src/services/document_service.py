"""Document download and text extraction service."""

import io
import logging
import re
import zipfile
from pathlib import Path
from urllib.parse import urljoin, urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from src.config import settings
from src.models import Tender, TenderStatus
from datetime import datetime

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for downloading and processing tender documents.

    Downloads documents from zakupki.gov.ru, extracts text,
    and stores it in the database (no S3 storage).
    """

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".rtf"}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    def __init__(self):
        self.base_url = "https://zakupki.gov.ru"
        self.proxy_url = settings.proxy_url

    def _get_http_client(self, timeout: float = 30.0) -> httpx.AsyncClient:
        """Get HTTP client with proxy and browser-like headers."""
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
            "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
            "Upgrade-Insecure-Requests": "1",
        }
        kwargs = {
            "timeout": timeout,
            "follow_redirects": True,
            "headers": headers,
        }
        if self.proxy_url:
            kwargs["proxy"] = self.proxy_url
        return httpx.AsyncClient(**kwargs)

    async def _fetch_tender_page(self, url: str, retries: int = 2) -> str | None:
        """Fetch tender page HTML with retry logic."""
        import asyncio

        for attempt in range(retries + 1):
            try:
                async with self._get_http_client(timeout=30.0) as client:
                    response = await client.get(url)
                    response.raise_for_status()
                    return response.text
            except httpx.HTTPStatusError as e:
                if e.response.status_code == 404:
                    # 404 is expected for cancelled/deleted tenders, don't retry
                    logger.warning(f"Tender page not found (404): {url}")
                    return None
                elif e.response.status_code >= 500 and attempt < retries:
                    # Server error, retry after delay
                    logger.warning(f"Server error {e.response.status_code}, retrying...")
                    await asyncio.sleep(2 ** attempt)
                    continue
                else:
                    logger.error(f"Error fetching tender page {url}: {e}")
                    return None
            except httpx.TimeoutException:
                if attempt < retries:
                    logger.warning(f"Timeout fetching {url}, retrying...")
                    await asyncio.sleep(2 ** attempt)
                    continue
                logger.error(f"Timeout fetching tender page after {retries} retries: {url}")
                return None
            except Exception as e:
                logger.error(f"Error fetching tender page {url}: {e}")
                return None

        return None

    def _extract_document_links(self, html: str, base_url: str) -> list[str]:
        """Extract document download links from tender page.

        zakupki.gov.ru document links patterns:
        - /epz/main/public/download/downloadDocument.html?id=...
        - Direct file links (.pdf, .docx, .doc, .zip, .rar, .xlsx, .xls)
        - /epz/order/notice/.../download/...
        - printForm links for generating documents
        """
        soup = BeautifulSoup(html, "lxml")
        links = []

        # Pattern 1: Look for document download links in zakupki.gov.ru format
        for a_tag in soup.find_all("a", href=True):
            href = a_tag["href"]
            href_lower = href.lower()

            # zakupki.gov.ru specific download patterns
            if any(pattern in href_lower for pattern in [
                "downloadDocument.html",
                "download/downloadDocument",
                "/download/",
                "fileDownload",
                "printForm",
                "getFile",
            ]):
                full_url = urljoin(base_url, href)
                links.append(full_url)
                continue

            # Direct file extensions
            if any(ext in href_lower for ext in [
                ".pdf", ".docx", ".doc", ".zip", ".rar",
                ".xlsx", ".xls", ".xml", ".sig"
            ]):
                full_url = urljoin(base_url, href)
                links.append(full_url)
                continue

        # Pattern 2: Look for attachment blocks (common in zakupki.gov.ru)
        # These may have specific classes like 'attachment', 'file-row', etc.
        attachment_containers = soup.find_all(
            class_=lambda x: x and any(
                cls in str(x).lower()
                for cls in ["attachment", "file", "document", "download"]
            )
        )
        for container in attachment_containers:
            for a_tag in container.find_all("a", href=True):
                href = a_tag["href"]
                if href and not href.startswith("#") and not href.startswith("javascript"):
                    full_url = urljoin(base_url, href)
                    if full_url not in links:
                        links.append(full_url)

        # Pattern 3: Look for onclick handlers that might contain download URLs
        for element in soup.find_all(attrs={"onclick": True}):
            onclick = element.get("onclick", "")
            # Extract URLs from onclick handlers
            url_match = re.search(r"['\"]([^'\"]*(?:download|file|document)[^'\"]*)['\"]", onclick, re.IGNORECASE)
            if url_match:
                href = url_match.group(1)
                if href and not href.startswith("javascript"):
                    full_url = urljoin(base_url, href)
                    if full_url not in links:
                        links.append(full_url)

        # Remove duplicates while preserving order
        seen = set()
        unique_links = []
        for link in links:
            if link not in seen:
                seen.add(link)
                unique_links.append(link)

        logger.info(f"Found {len(unique_links)} document links")
        return unique_links

    def _extract_deadline(self, html: str) -> datetime | None:
        """Extract application deadline from tender common-info page.

        zakupki.gov.ru deadline patterns:
        - "Дата и время окончания подачи заявок" / "Дата и время окончания срока подачи заявок"
        - "Дата окончания подачи заявок"
        - Format: DD.MM.YYYY HH:MM (Moscow time)
        """
        soup = BeautifulSoup(html, "lxml")

        # Pattern labels to look for (case insensitive)
        deadline_labels = [
            "дата и время окончания подачи заявок",
            "дата и время окончания срока подачи заявок",
            "дата окончания подачи заявок",
            "окончание подачи заявок",
            "срок подачи заявок",
            "дата и время окончания срока подачи",
        ]

        # Pattern 1: section__title + section__info structure (most common on zakupki.gov.ru)
        for section in soup.find_all("section", class_="blockInfo__section"):
            title_elem = section.find("span", class_="section__title")
            info_elem = section.find("span", class_="section__info")
            if title_elem and info_elem:
                label_text = title_elem.get_text(strip=True).lower()
                for pattern in deadline_labels:
                    if pattern in label_text:
                        value_text = info_elem.get_text(strip=True)
                        deadline = self._parse_deadline_string(value_text)
                        if deadline:
                            logger.info(f"Deadline extracted from section: {deadline}")
                            return deadline

        # Pattern 2: Try to find in table structure (common format)
        for row in soup.find_all("tr"):
            cells = row.find_all(["td", "th"])
            if len(cells) >= 2:
                label_text = cells[0].get_text(strip=True).lower()
                for pattern in deadline_labels:
                    if pattern in label_text:
                        value_text = cells[1].get_text(strip=True)
                        deadline = self._parse_deadline_string(value_text)
                        if deadline:
                            logger.info(f"Deadline extracted from table: {deadline}")
                            return deadline

        # Pattern 3: div/span based layout (older format)
        for label in deadline_labels:
            # Find by text content
            elements = soup.find_all(string=lambda t: t and label in t.lower())
            for elem in elements:
                # Look for adjacent value element
                parent = elem.parent
                if parent:
                    # Check next sibling
                    next_elem = parent.find_next_sibling()
                    if next_elem:
                        value_text = next_elem.get_text(strip=True)
                        deadline = self._parse_deadline_string(value_text)
                        if deadline:
                            logger.info(f"Deadline extracted from div: {deadline}")
                            return deadline

                    # Check parent's next sibling
                    grandparent = parent.parent
                    if grandparent:
                        next_elem = grandparent.find_next_sibling()
                        if next_elem:
                            value_text = next_elem.get_text(strip=True)
                            deadline = self._parse_deadline_string(value_text)
                            if deadline:
                                logger.info(f"Deadline extracted from grandparent: {deadline}")
                                return deadline

        # Fallback: regex search in full HTML
        deadline_pattern = r'(?:окончани[яе].*?подачи.*?заявок|срок.*?подачи)[^\d]*(\d{2}\.\d{2}\.\d{4})\s*(?:(\d{2}:\d{2}))?'
        match = re.search(deadline_pattern, html, re.IGNORECASE | re.DOTALL)
        if match:
            date_str = match.group(1)
            time_str = match.group(2) if match.group(2) else "23:59"
            deadline = self._parse_deadline_string(f"{date_str} {time_str}")
            if deadline:
                logger.info(f"Deadline extracted from regex: {deadline}")
                return deadline

        logger.debug("Deadline not found in HTML")
        return None

    def _parse_deadline_string(self, text: str) -> datetime | None:
        """Parse deadline string into datetime.

        Formats:
        - DD.MM.YYYY HH:MM
        - DD.MM.YYYY
        - DD.MM.YYYY HH:MM:SS
        """
        if not text:
            return None

        # Clean up text
        text = text.strip()

        # Extract date and time using regex
        match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})\s*(\d{2}:\d{2})?(?::\d{2})?', text)
        if not match:
            return None

        try:
            day = int(match.group(1))
            month = int(match.group(2))
            year = int(match.group(3))
            time_part = match.group(4)

            if time_part:
                hour, minute = map(int, time_part.split(":"))
            else:
                hour, minute = 23, 59  # Default to end of day

            # Create datetime (Moscow time, but store as naive for simplicity)
            return datetime(year, month, day, hour, minute)
        except (ValueError, TypeError) as e:
            logger.debug(f"Failed to parse deadline '{text}': {e}")
            return None

    async def _download_file(self, url: str) -> tuple[bytes, str] | None:
        """Download file and return content with filename."""
        try:
            async with self._get_http_client(timeout=60.0) as client:
                response = await client.get(url)
                response.raise_for_status()

                # Check file size
                content = response.content
                if len(content) > self.MAX_FILE_SIZE:
                    logger.warning(f"File too large: {url}")
                    return None

                # Check if we got HTML instead of a document (common for printForm errors)
                content_type = response.headers.get("content-type", "")
                if "text/html" in content_type and len(content) < 10000:
                    # Likely an error page, not a real document
                    logger.debug(f"Got HTML response instead of document: {url}")
                    return None

                # Extract filename from URL or Content-Disposition
                filename = Path(urlparse(url).path).name
                if "content-disposition" in response.headers:
                    cd = response.headers["content-disposition"]
                    if "filename=" in cd:
                        filename = re.findall(r'filename="?([^";\n]+)"?', cd)[0]

                # Default filename for printForm
                if not filename or filename == "view.html":
                    filename = f"document_{hash(url) % 10000}.pdf"

                return content, filename

        except httpx.HTTPStatusError as e:
            # 404 for printForm is common and expected, log at debug level
            if e.response.status_code == 404 and "printForm" in url:
                logger.debug(f"PrintForm not available (404): {url}")
            else:
                logger.error(f"Error downloading file {url}: {e}")
            return None
        except Exception as e:
            logger.error(f"Error downloading file {url}: {e}")
            return None

    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""

    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def _extract_text_from_zip(self, content: bytes) -> str:
        """Extract text from files inside ZIP archive."""
        texts = []
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                for filename in zf.namelist():
                    ext = Path(filename).suffix.lower()
                    if ext in self.SUPPORTED_EXTENSIONS:
                        try:
                            file_content = zf.read(filename)
                            if ext == ".pdf":
                                texts.append(self._extract_text_from_pdf(file_content))
                            elif ext == ".docx":
                                texts.append(self._extract_text_from_docx(file_content))
                            elif ext == ".txt":
                                texts.append(file_content.decode("utf-8", errors="ignore"))
                        except Exception as e:
                            logger.warning(f"Error processing {filename} in ZIP: {e}")
        except Exception as e:
            logger.error(f"Error processing ZIP file: {e}")

        return "\n\n".join(filter(None, texts))

    def _extract_text(self, content: bytes, filename: str) -> str:
        """Extract text based on file type."""
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._extract_text_from_pdf(content)
        elif ext == ".docx":
            return self._extract_text_from_docx(content)
        elif ext == ".zip":
            return self._extract_text_from_zip(content)
        elif ext == ".txt":
            return content.decode("utf-8", errors="ignore")

        return ""

    # Known tender type URL patterns on zakupki.gov.ru
    # Format: /epz/order/notice/{type}/view/{page}.html
    TENDER_TYPES = [
        "ea20",    # Electronic auction (44-FZ)
        "zk20",    # Request for quotations (44-FZ)
        "ep44",    # Electronic procedure (44-FZ)
        "ok20",    # Open competition (44-FZ)
        "zkk20",   # Competitive request for quotations
        "oku20",   # Open competition in electronic form
        "ezk20",   # Electronic request for quotations
        "ezp20",   # Electronic request for proposals
        "ek20",    # Electronic competition
        "ea44",    # Electronic auction 44-FZ
        "ok44",    # Open competition 44-FZ
        "zp44",    # Request for proposals 44-FZ
    ]

    def _build_documents_url(self, tender_url: str) -> str:
        """Convert tender URL to documents page URL.

        zakupki.gov.ru tender pages have multiple views:
        - common-info.html - general information
        - documents.html - documents and attachments
        - supplier-results.html - supplier info

        RSS feed gives common-info URLs, but documents are on documents.html
        """
        # Replace common-info.html with documents.html
        documents_url = tender_url.replace("common-info.html", "documents.html")

        # Also handle other view pages
        documents_url = documents_url.replace("supplier-results.html", "documents.html")
        documents_url = documents_url.replace("event-journal.html", "documents.html")

        return documents_url

    def _get_alternative_urls(self, tender_url: str, reg_number: str) -> list[str]:
        """Generate alternative URLs to try if the main URL fails.

        Sometimes the tender type in the URL is incorrect,
        try common types as alternatives.
        """
        alternatives = []

        # Extract base parts
        for tender_type in self.TENDER_TYPES:
            alt_url = f"{self.base_url}/epz/order/notice/{tender_type}/view/documents.html?regNumber={reg_number}"
            if alt_url != tender_url and alt_url not in alternatives:
                alternatives.append(alt_url)

        return alternatives[:5]  # Limit to 5 alternatives

    async def process_tender(self, tender: Tender, db: AsyncSession) -> bool:
        """
        Download documents for a tender and extract text.

        Text is stored directly in the database (raw_text field).
        No S3 storage is used.

        Args:
            tender: Tender entity to process
            db: Database session

        Returns:
            True if successful, False otherwise
        """
        logger.info(f"Processing tender {tender.external_id}")

        # Build full URL if relative
        tender_url = tender.url
        if not tender_url.startswith("http"):
            tender_url = urljoin(self.base_url, tender_url)

        # First, fetch common-info page to extract deadline
        common_info_url = tender_url
        if "documents.html" in common_info_url:
            common_info_url = common_info_url.replace("documents.html", "common-info.html")
        elif "supplier-results.html" in common_info_url:
            common_info_url = common_info_url.replace("supplier-results.html", "common-info.html")

        logger.debug(f"Common info URL: {common_info_url}")
        common_info_html = await self._fetch_tender_page(common_info_url)

        if common_info_html:
            # Extract deadline from common-info page
            deadline = self._extract_deadline(common_info_html)
            if deadline:
                tender.deadline = deadline
                logger.info(f"Tender {tender.external_id} deadline: {deadline}")

        # Convert to documents page URL
        documents_url = self._build_documents_url(tender_url)
        logger.debug(f"Documents URL: {documents_url}")

        # Fetch documents page (not common-info)
        html = await self._fetch_tender_page(documents_url)

        if not html:
            # Try alternative URL patterns with different tender types
            logger.info(f"Primary URL failed, trying alternatives for {tender.external_id}")
            alternatives = self._get_alternative_urls(documents_url, tender.external_id)
            for alt_url in alternatives:
                html = await self._fetch_tender_page(alt_url)
                if html:
                    logger.info(f"Found working alternative URL: {alt_url}")
                    documents_url = alt_url
                    break

        if not html:
            # Last resort: try original URL
            logger.warning(f"All alternatives failed, trying original URL")
            html = await self._fetch_tender_page(tender_url)

        if not html:
            logger.error(f"Could not fetch tender page: {tender.url}")
            # Mark as downloaded with empty text to avoid retrying indefinitely
            tender.status = TenderStatus.DOWNLOADED
            tender.raw_text = ""
            await db.commit()
            return False

        # Extract document links
        doc_links = self._extract_document_links(html, documents_url)
        if not doc_links:
            logger.warning(f"No document links found for tender {tender.external_id}")
            # Still mark as downloaded with empty text
            tender.status = TenderStatus.DOWNLOADED
            tender.raw_text = ""
            await db.commit()
            return True

        # Download and extract text from all documents
        all_texts = []

        for url in doc_links[:10]:  # Limit to 10 files
            result = await self._download_file(url)
            if not result:
                continue

            content, filename = result

            # Extract text (no S3 upload - text only)
            text = self._extract_text(content, filename)
            if text:
                all_texts.append(f"--- {filename} ---\n{text}")

        # Update tender with extracted text
        tender.raw_text = "\n\n".join(all_texts) if all_texts else ""
        tender.status = TenderStatus.DOWNLOADED

        await db.commit()
        logger.info(
            f"Tender {tender.external_id} processed: "
            f"{len(all_texts)} documents, {len(tender.raw_text)} chars"
        )
        return True

    async def process_new_tenders(self, db: AsyncSession, limit: int = 5) -> int:
        """
        Process tenders with NEW status.

        Args:
            db: Database session
            limit: Maximum number of tenders to process

        Returns:
            Number of tenders processed
        """
        # Get NEW tenders
        result = await db.execute(
            select(Tender)
            .where(Tender.status == TenderStatus.NEW)
            .order_by(Tender.created_at)
            .limit(limit)
        )
        tenders = result.scalars().all()

        processed = 0
        for tender in tenders:
            try:
                if await self.process_tender(tender, db):
                    processed += 1
            except Exception as e:
                logger.error(f"Error processing tender {tender.external_id}: {e}")

        return processed


# Singleton instance
document_service = DocumentService()
